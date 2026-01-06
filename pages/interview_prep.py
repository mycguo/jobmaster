"""
Interview Preparation Dashboard

Main hub for managing interview prep materials:
- Questions and answers
- Technical concepts
- Company research
- Practice sessions
"""

import streamlit as st
import sys
from datetime import datetime
import random

# Add parent directory to path
sys.path.insert(0, '.')

from storage.interview_db import InterviewDB
from storage.auth_utils import (
    is_user_logged_in,
    logout,
    render_login_button,
    render_linkedin_login_button,
    handle_linkedin_callback,
    is_linkedin_configured
)
from models.interview_prep import (
    create_interview_question,
    create_technical_concept,
    create_company_research,
    create_practice_session,
    PracticeSession
)
from storage.pg_vector_store import PgVectorStore
from pages.upload_docs import get_text_chunks
import io
from langchain_google_genai import ChatGoogleGenerativeAI
import json
from PyPDF2 import PdfReader
import docx


def extract_text_from_file(file_bytes, filename):
    """Extract text content from uploaded file (PDF, Word, or Text)"""
    try:
        file_lower = filename.lower()

        if file_lower.endswith('.pdf'):
            # Extract from PDF
            pdf = PdfReader(io.BytesIO(file_bytes))
            text = ""
            metadata = {
                'filename': filename,
                'num_pages': len(pdf.pages),
                'type': 'pdf'
            }

            for page in pdf.pages:
                text += page.extract_text()

            return text, metadata

        elif file_lower.endswith('.docx'):
            # Extract from Word document
            doc = docx.Document(io.BytesIO(file_bytes))
            text = ""
            metadata = {
                'filename': filename,
                'num_paragraphs': len(doc.paragraphs),
                'type': 'docx'
            }

            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"

            return text, metadata

        elif file_lower.endswith('.txt'):
            # Extract from text file
            text = file_bytes.decode('utf-8')
            metadata = {
                'filename': filename,
                'type': 'txt'
            }

            return text, metadata

        else:
            return None, {'error': 'Unsupported file type'}

    except Exception as e:
        return None, {'error': str(e)}


def add_question_to_vector_store(question, answer, metadata):
    """Add question and answer to vector store for semantic search"""
    try:
        vector_store = PgVectorStore()

        # Create searchable content
        content = f"""Interview Question: {question}

Answer: {answer}

Type: {metadata.get('type', '')}
Category: {metadata.get('category', '')}
Companies: {', '.join(metadata.get('companies', []))}
Tags: {', '.join(metadata.get('tags', []))}"""

        # Add to vector store
        text_chunks = get_text_chunks(content)
        metadatas = [{
            'source': 'interview_question',
            'question_id': metadata.get('question_id'),
            'type': 'interview_prep',
            'timestamp': datetime.now().isoformat(),
            **metadata
        } for _ in text_chunks]

        vector_store.add_texts(text_chunks, metadatas=metadatas)
        return True
    except Exception as e:
        st.error(f"Error adding to vector store: {str(e)}")
        return False


def add_document_to_vector_store(file_content, filename, metadata):
    """Add entire document to vector store for semantic search"""
    try:
        vector_store = PgVectorStore()

        # Add document metadata
        content = f"""Interview Prep Document: {filename}

{file_content}

Category: {metadata.get('category', 'interview_prep')}
Tags: {', '.join(metadata.get('tags', []))}"""

        # Chunk and add to vector store
        text_chunks = get_text_chunks(content)
        metadatas = [{
            'source': 'interview_prep_document',
            'filename': filename,
            'type': 'interview_prep',
            'timestamp': datetime.now().isoformat(),
            **metadata
        } for _ in text_chunks]

        vector_store.add_texts(text_chunks, metadatas=metadatas)
        return True, len(text_chunks)
    except Exception as e:
        return False, str(e)


def parse_questions_from_document(text_content):
    """Use AI to parse questions and answers from document"""
    try:
        model = ChatGoogleGenerativeAI(model="gemini-2.5-flash", temperature=0.0)

        prompt = f"""Parse this interview preparation document and extract all questions and answers.

Document:
{text_content}

Please extract all interview questions and their answers. Return a JSON array with this structure:
[
  {{
    "question": "the question text",
    "answer": "the answer text",
    "type": "behavioral or technical or system-design or case-study",
    "category": "best guess category (leadership, algorithms, conflict, etc.)"
  }}
]

Only return valid JSON, no other text. If no clear Q&A found, return empty array []."""

        response = model.invoke(prompt)

        # Parse JSON response
        content = response.content.strip()
        # Remove markdown code blocks if present
        if content.startswith("```"):
            content = content.split("```")[1]
            if content.startswith("json"):
                content = content[4:]

        parsed = json.loads(content.strip())
        return True, parsed
    except Exception as e:
        return False, f"Error parsing document: {str(e)}"


def show_upload_document_form(db: InterviewDB):
    """Show form to upload interview prep document"""
    st.subheader("📄 Upload Interview Prep Document")

    st.markdown("""
    Upload a document containing your interview questions and answers.
    Supports: PDF, Word, Text files.
    """)

    with st.form("upload_document_form", clear_on_submit=True):
        uploaded_file = st.file_uploader(
            "Choose a file",
            type=['pdf', 'docx', 'txt'],
            help="Upload your interview prep notes, Q&A document, etc."
        )

        col1, col2 = st.columns(2)

        with col1:
            category = st.text_input(
                "Category (optional)",
                placeholder="e.g., behavioral, technical",
                help="Category for organization"
            )

        with col2:
            tags = st.text_input(
                "Tags (comma-separated, optional)",
                placeholder="e.g., leadership, algorithms",
                help="Tags for easier searching"
            )

        st.markdown("### Processing Options")

        process_mode = st.radio(
            "How do you want to process this document?",
            [
                "📦 Quick: Just store in searchable knowledge base (Recommended)",
                "🤖 Smart: AI parse into individual questions (slower)"
            ],
            help="Quick mode stores the whole document. Smart mode uses AI to extract Q&A pairs."
        )

        submit = st.form_submit_button("Upload Document", type="primary")

        if submit and uploaded_file:
            with st.spinner("Processing document..."):
                # Extract text from file
                file_bytes = uploaded_file.read()
                text_content, metadata = extract_text_from_file(file_bytes, uploaded_file.name)

                if not text_content:
                    st.error("Could not extract text from file")
                    return

                # Prepare metadata
                doc_metadata = {
                    'category': category if category else 'interview_prep',
                    'tags': [t.strip() for t in tags.split(',')] if tags else [],
                    'original_filename': uploaded_file.name
                }

                if "Quick" in process_mode:
                    # Quick mode: Just add to vector store
                    success, result = add_document_to_vector_store(
                        text_content,
                        uploaded_file.name,
                        doc_metadata
                    )

                    if success:
                        st.success(f"✅ Document uploaded! Added {result} chunks to knowledge base.")
                        st.info("💡 You can now ask questions about this document in the chat!")

                        with st.expander("📄 Document Preview"):
                            st.text(text_content[:1000] + "..." if len(text_content) > 1000 else text_content)
                    else:
                        st.error(f"❌ Error: {result}")

                else:
                    # Smart mode: Parse questions
                    st.info("🤖 Parsing questions with AI... This may take 10-20 seconds.")
                    success, result = parse_questions_from_document(text_content[:10000])  # Limit for API

                    if not success:
                        st.error(result)
                        st.info("Falling back to quick mode...")
                        success, chunks = add_document_to_vector_store(
                            text_content,
                            uploaded_file.name,
                            doc_metadata
                        )
                        if success:
                            st.success(f"✅ Document stored in knowledge base ({chunks} chunks)")
                        return

                    if not result or len(result) == 0:
                        st.warning("No questions found. Storing document as-is.")
                        success, chunks = add_document_to_vector_store(
                            text_content,
                            uploaded_file.name,
                            doc_metadata
                        )
                        if success:
                            st.success(f"✅ Document stored in knowledge base ({chunks} chunks)")
                        return

                    # Store parsed questions in session state for review
                    st.session_state['parsed_questions'] = result
                    st.session_state['parsed_metadata'] = doc_metadata
                    st.success(f"✅ Found {len(result)} questions! Review them below.")
                    st.rerun()


def show_parsed_questions_review(db: InterviewDB):
    """Show review interface for parsed questions (outside form)"""
    if 'parsed_questions' not in st.session_state:
        return

    result = st.session_state['parsed_questions']
    doc_metadata = st.session_state.get('parsed_metadata', {})

    st.subheader(f"📋 Review {len(result)} Parsed Questions")

    with st.expander(f"Review and Select Questions to Save", expanded=True):
        questions_to_save = st.multiselect(
            "Select questions to save:",
            range(len(result)),
            default=list(range(len(result))),
            format_func=lambda i: f"Q{i+1}: {result[i]['question'][:60]}..."
        )

        for i in questions_to_save:
            q_data = result[i]
            st.markdown(f"**Q{i+1}:** {q_data['question']}")
            st.caption(f"Type: {q_data.get('type', 'unknown')} | Category: {q_data.get('category', 'unknown')}")
            with st.expander("View Answer"):
                st.write(q_data['answer'])
            st.divider()

    col1, col2 = st.columns([1, 4])

    with col1:
        if st.button("💾 Save Selected Questions", key="save_parsed", type="primary"):
            saved_count = 0
            for i in questions_to_save:
                q_data = result[i]

                # Create question
                q = create_interview_question(
                    question=q_data['question'],
                    type=q_data.get('type', 'technical'),
                    category=q_data.get('category', doc_metadata.get('category', 'general')),
                    difficulty='medium',  # Default
                    answer_full=q_data['answer'],
                    tags=doc_metadata.get('tags', [])
                )

                # Save to database
                q_id = db.add_question(q)

                # Add to vector store
                add_question_to_vector_store(
                    question=q_data['question'],
                    answer=q_data['answer'],
                    metadata={
                        'question_id': q_id,
                        'type': q_data.get('type', 'technical'),
                        'category': q_data.get('category', 'general'),
                        'tags': doc_metadata.get('tags', [])
                    }
                )
                saved_count += 1

            st.success(f"✅ Saved {saved_count} questions!")
            # Clear session state
            del st.session_state['parsed_questions']
            del st.session_state['parsed_metadata']
            st.balloons()
            st.rerun()

    with col2:
        if st.button("✕ Cancel", key="cancel_parsed"):
            del st.session_state['parsed_questions']
            del st.session_state['parsed_metadata']
            st.rerun()


def show_add_question_form(db: InterviewDB):
    """Show form to add new interview question"""
    st.subheader("➕ Add New Question & Answer")

    with st.form("add_question_form", clear_on_submit=True):
        # Question details
        question = st.text_area(
            "Interview Question *",
            placeholder="e.g., Tell me about a time you led a difficult project",
            height=100,
            help="The interview question you want to prepare for"
        )

        col1, col2, col3 = st.columns(3)

        with col1:
            q_type = st.selectbox(
                "Type *",
                ["behavioral", "technical", "system-design", "case-study"],
                help="Type of interview question"
            )

        with col2:
            category = st.text_input(
                "Category *",
                placeholder="e.g., leadership, algorithms",
                help="Category for organization (leadership, conflict, algorithms, etc.)"
            )

        with col3:
            difficulty = st.selectbox(
                "Difficulty",
                ["easy", "medium", "hard"]
            )

        # Answer section
        st.markdown("### Your Answer")

        # Choose answer format
        answer_format = st.radio(
            "Answer Format",
            ["Full Answer", "STAR Format (Behavioral)"],
            help="STAR: Situation, Task, Action, Result"
        )

        if answer_format == "STAR Format (Behavioral)":
            situation = st.text_area(
                "Situation",
                placeholder="Describe the context and background...",
                height=80
            )
            task = st.text_area(
                "Task",
                placeholder="What was the challenge or goal?...",
                height=80
            )
            action = st.text_area(
                "Action",
                placeholder="What specific actions did you take?...",
                height=80
            )
            result = st.text_area(
                "Result",
                placeholder="What was the outcome? Include metrics if possible...",
                height=80
            )

            # Combine for full answer
            answer_full = f"""**Situation:** {situation}

**Task:** {task}

**Action:** {action}

**Result:** {result}"""

            answer_star = {
                'situation': situation,
                'task': task,
                'action': action,
                'result': result
            } if all([situation, task, action, result]) else None
        else:
            answer_full = st.text_area(
                "Your Answer *",
                placeholder="Write your complete answer here...",
                height=200
            )
            answer_star = None

        # Additional details
        tags = st.text_input(
            "Tags (comma-separated)",
            placeholder="e.g., leadership, team-management",
            help="Tags for easier searching"
        )

        notes = st.text_area(
            "Notes (optional)",
            placeholder="Additional notes, variations, tips...",
            height=80
        )

        confidence = st.slider(
            "Confidence Level",
            1, 5, 3,
            help="How confident are you in this answer? (1=Low, 5=High)"
        )

        # Vector store option
        add_to_vector = st.checkbox(
            "💾 Add to searchable knowledge base",
            value=True,
            help="Store in vector DB for semantic search and RAG queries"
        )

        submit = st.form_submit_button("Add Question", type="primary")

        if submit:
            if not question or not category or not answer_full:
                st.error("⚠️ Please fill in all required fields!")
            else:
                try:
                    # Parse tags (companies removed from UI)
                    company_list = []  # Keep empty for backward compatibility
                    tag_list = [t.strip() for t in tags.split(',')] if tags else []

                    # Create question
                    q = create_interview_question(
                        question=question,
                        type=q_type,
                        category=category,
                        difficulty=difficulty,
                        answer_full=answer_full,
                        answer_star=answer_star,
                        notes=notes,
                        tags=tag_list,
                        companies=company_list,
                        confidence_level=confidence
                    )

                    # Save to database
                    q_id = db.add_question(q)
                    st.success(f"✅ Question added successfully! (ID: {q_id})")

                    # Add to vector store if requested
                    if add_to_vector:
                        with st.spinner("Adding to knowledge base..."):
                            success = add_question_to_vector_store(
                                question=question,
                                answer=answer_full,
                                metadata={
                                    'question_id': q_id,
                                    'type': q_type,
                                    'category': category,
                                    'companies': company_list,
                                    'tags': tag_list
                                }
                            )
                            if success:
                                st.success("✅ Added to searchable knowledge base!")

                    st.balloons()
                    st.rerun()

                except Exception as e:
                    st.error(f"Error: {str(e)}")


def show_practice_mode(db: InterviewDB):
    """Practice Mode - Interactive question practice session"""
    st.subheader("🎓 Practice Mode")
    st.markdown("Practice your interview questions with timer and self-evaluation")

    # Initialize practice session state
    if 'practice_session' not in st.session_state:
        st.session_state['practice_session'] = None
        st.session_state['current_question_index'] = 0
        st.session_state['practice_questions'] = []
        st.session_state['start_time'] = None
        st.session_state['show_answer'] = False

    # Get all questions
    all_questions = db.list_questions()

    if not all_questions:
        st.warning("No questions available. Add some questions first!")
        return

    # Practice session configuration
    if st.session_state['practice_session'] is None:
        st.markdown("### Configure Practice Session")

        col1, col2 = st.columns(2)

        with col1:
            # Filter options
            filter_type = st.selectbox(
                "Question Type",
                ["All", "behavioral", "technical", "system-design", "case-study"],
                help="Filter questions by type"
            )

            filter_difficulty = st.selectbox(
                "Difficulty",
                ["All", "easy", "medium", "hard"],
                help="Filter questions by difficulty"
            )

        with col2:
            num_questions = st.slider(
                "Number of Questions",
                min_value=1,
                max_value=min(20, len(all_questions)),
                value=min(5, len(all_questions)),
                help="How many questions to practice"
            )

            practice_mode = st.radio(
                "Practice Mode",
                ["Random", "Least Practiced", "Low Confidence"],
                help="How to select questions"
            )

        if st.button("🎯 Start Practice Session", type="primary", width="stretch"):
            # Filter questions
            filtered_questions = all_questions

            if filter_type != "All":
                filtered_questions = [q for q in filtered_questions if q.type == filter_type]

            if filter_difficulty != "All":
                filtered_questions = [q for q in filtered_questions if q.difficulty == filter_difficulty]

            if not filtered_questions:
                st.error("No questions match your filters!")
                return

            # Select questions based on mode
            if practice_mode == "Random":
                selected = random.sample(filtered_questions, min(num_questions, len(filtered_questions)))
            elif practice_mode == "Least Practiced":
                sorted_q = sorted(filtered_questions, key=lambda x: x.practice_count)
                selected = sorted_q[:num_questions]
            else:  # Low Confidence
                sorted_q = sorted(filtered_questions, key=lambda x: x.confidence_level)
                selected = sorted_q[:num_questions]

            # Create practice session
            session = create_practice_session(
                session_type=filter_type if filter_type != "All" else "general",
                duration_minutes=0
            )

            st.session_state['practice_session'] = session
            st.session_state['practice_questions'] = selected
            st.session_state['current_question_index'] = 0
            st.session_state['start_time'] = datetime.now()
            st.session_state['show_answer'] = False
            st.rerun()

    else:
        # Active practice session
        session = st.session_state['practice_session']
        questions = st.session_state['practice_questions']
        current_index = st.session_state['current_question_index']

        if current_index >= len(questions):
            # Session complete
            st.success("🎉 Practice Session Complete!")

            elapsed = (datetime.now() - st.session_state['start_time']).seconds // 60
            session.duration_minutes = elapsed

            # Save session
            db.add_practice_session(session)

            st.markdown(f"### Session Summary")
            col1, col2, col3 = st.columns(3)

            with col1:
                st.metric("Questions Practiced", len(questions))
            with col2:
                st.metric("Duration", f"{elapsed} min")
            with col3:
                avg_rating = session.get_average_rating()
                st.metric("Avg Performance", f"{avg_rating:.1f}/5")

            if st.button("✨ Start New Session"):
                st.session_state['practice_session'] = None
                st.session_state['current_question_index'] = 0
                st.session_state['practice_questions'] = []
                st.session_state['start_time'] = None
                st.session_state['show_answer'] = False
                st.rerun()

            return

        # Show current question
        current_q = questions[current_index]

        st.progress((current_index + 1) / len(questions))
        st.caption(f"Question {current_index + 1} of {len(questions)}")

        st.markdown("---")

        # Question details
        col1, col2, col3 = st.columns(3)
        with col1:
            st.write(f"**Type:** {current_q.get_display_type()}")
        with col2:
            st.write(f"**Difficulty:** {current_q.get_difficulty_emoji()} {current_q.difficulty.title()}")
        with col3:
            st.write(f"**Category:** {current_q.category.title()}")

        st.markdown("### Question")
        st.markdown(f"**{current_q.question}**")

        # Timer
        elapsed_seconds = (datetime.now() - st.session_state['start_time']).seconds
        elapsed_minutes = elapsed_seconds // 60
        elapsed_secs = elapsed_seconds % 60
        st.caption(f"⏱️ Session Time: {elapsed_minutes:02d}:{elapsed_secs:02d}")

        st.markdown("---")

        # Show/Hide answer
        col1, col2 = st.columns([1, 3])

        with col1:
            if st.button("👁️ Show Answer" if not st.session_state['show_answer'] else "🙈 Hide Answer",
                        width="stretch"):
                st.session_state['show_answer'] = not st.session_state['show_answer']
                st.rerun()

        if st.session_state['show_answer']:
            st.markdown("### Answer")
            st.markdown(current_q.answer_full)

            if current_q.answer_star:
                with st.expander("📋 STAR Format"):
                    if 'situation' in current_q.answer_star:
                        st.markdown(f"**Situation:** {current_q.answer_star['situation']}")
                    if 'task' in current_q.answer_star:
                        st.markdown(f"**Task:** {current_q.answer_star['task']}")
                    if 'action' in current_q.answer_star:
                        st.markdown(f"**Action:** {current_q.answer_star['action']}")
                    if 'result' in current_q.answer_star:
                        st.markdown(f"**Result:** {current_q.answer_star['result']}")

            if current_q.notes:
                with st.expander("📝 Notes"):
                    st.markdown(current_q.notes)

            st.markdown("---")

            # Self-evaluation
            st.markdown("### Self-Evaluation")

            col1, col2 = st.columns(2)

            with col1:
                performance_rating = st.slider(
                    "How well did you answer?",
                    min_value=1,
                    max_value=5,
                    value=3,
                    help="1 = Poor, 5 = Excellent"
                )

            with col2:
                new_confidence = st.slider(
                    "Confidence Level",
                    min_value=1,
                    max_value=5,
                    value=current_q.confidence_level,
                    help="Update your confidence for this question"
                )

            performance_notes = st.text_area(
                "Notes (optional)",
                placeholder="What went well? What needs improvement?",
                height=100
            )

            if st.button("➡️ Next Question", type="primary", width="stretch"):
                # Record performance
                session.add_question(current_q.id, performance_rating, performance_notes)

                # Update question
                current_q.mark_practiced()
                current_q.update_confidence(new_confidence)
                db.update_question(current_q)

                # Move to next
                st.session_state['current_question_index'] += 1
                st.session_state['show_answer'] = False
                st.rerun()


def show_question_edit_form(db: InterviewDB, question, question_id: str, edit_mode_key: str):
    """Show comprehensive edit form for all question fields"""
    st.markdown("### ✏️ Edit Question")

    with st.form("edit_question_form"):
        # Question Text
        st.markdown("**Question**")
        question_text = st.text_area(
            "Question",
            value=question.question,
            height=100,
            label_visibility="collapsed",
            help="The interview question"
        )

        # Type, Category, Difficulty in columns
        col1, col2, col3 = st.columns(3)

        with col1:
            question_type = st.selectbox(
                "Type",
                ["behavioral", "technical", "system-design", "case-study"],
                index=["behavioral", "technical", "system-design", "case-study"].index(question.type)
            )

        with col2:
            category = st.text_input(
                "Category",
                value=question.category,
                placeholder="e.g., leadership, algorithms"
            )

        with col3:
            difficulty = st.selectbox(
                "Difficulty",
                ["easy", "medium", "hard"],
                index=["easy", "medium", "hard"].index(question.difficulty)
            )

        # Importance and Confidence
        col1, col2 = st.columns(2)

        with col1:
            importance = st.slider(
                "Importance",
                min_value=1,
                max_value=10,
                value=question.importance,
                help="How important is this question? (1-10)"
            )

        with col2:
            confidence_level = st.slider(
                "Confidence Level",
                min_value=1,
                max_value=5,
                value=question.confidence_level,
                help="How confident are you with this answer? (1-5)"
            )

        # Tags and Companies
        col1, col2 = st.columns(2)

        with col1:
            tags_input = st.text_input(
                "Tags (comma-separated)",
                value=", ".join(question.tags) if question.tags else "",
                placeholder="e.g., leadership, teamwork"
            )

        with col2:
            companies_input = st.text_input(
                "Companies (comma-separated)",
                value=", ".join(question.companies) if question.companies else "",
                placeholder="e.g., Google, Amazon"
            )

        st.divider()

        # Answer
        st.markdown("**Answer**")
        answer_full = st.text_area(
            "Answer",
            value=question.answer_full,
            height=200,
            label_visibility="collapsed",
            help="Your complete answer to this question"
        )

        # STAR Format
        st.markdown("**STAR Format (Optional)**")

        col1, col2 = st.columns(2)

        with col1:
            star_situation = st.text_area(
                "Situation",
                value=question.answer_star.get('situation', '') if question.answer_star else '',
                height=80
            )

            star_task = st.text_area(
                "Task",
                value=question.answer_star.get('task', '') if question.answer_star else '',
                height=80
            )

        with col2:
            star_action = st.text_area(
                "Action",
                value=question.answer_star.get('action', '') if question.answer_star else '',
                height=80
            )

            star_result = st.text_area(
                "Result",
                value=question.answer_star.get('result', '') if question.answer_star else '',
                height=80
            )

        # Notes
        st.markdown("**Notes**")
        notes = st.text_area(
            "Notes",
            value=question.notes,
            height=100,
            label_visibility="collapsed",
            help="Additional notes about this question"
        )

        # Submit buttons
        col1, col2 = st.columns([1, 4])

        with col1:
            save = st.form_submit_button("💾 Save", type="primary", width="stretch")

        with col2:
            cancel = st.form_submit_button("✕ Cancel", width="stretch")

        if cancel:
            st.session_state[edit_mode_key] = False
            st.rerun()

        if save:
            if not question_text or not question_text.strip():
                st.error("Question text cannot be empty")
            elif not category or not category.strip():
                st.error("Category cannot be empty")
            else:
                # Parse tags and companies
                tags = [t.strip() for t in tags_input.split(",") if t.strip()]
                companies = [c.strip() for c in companies_input.split(",") if c.strip()]

                # Update question object
                question.question = question_text.strip()
                question.type = question_type
                question.category = category.strip()
                question.difficulty = difficulty
                question.importance = importance
                question.confidence_level = confidence_level
                question.tags = tags
                question.companies = companies
                question.answer_full = answer_full
                question.notes = notes

                # Update STAR format if any field is filled
                if star_situation or star_task or star_action or star_result:
                    question.answer_star = {
                        'situation': star_situation,
                        'task': star_task,
                        'action': star_action,
                        'result': star_result
                    }
                else:
                    question.answer_star = None

                # Update timestamp
                question.updated_at = datetime.now().isoformat()

                # Save to database
                db.update_question(question)

                # Exit edit mode
                st.session_state[edit_mode_key] = False

                st.success("✅ Question updated successfully!")
                st.rerun()


def show_question_detail(db: InterviewDB, question_id: str):
    """Show detailed view of a single question"""
    question = db.get_question(question_id)

    if not question:
        st.error("Question not found!")
        return

    # Check if in edit mode
    edit_mode_key = f'edit_question_detail_{question_id}'
    is_editing = st.session_state.get(edit_mode_key, False)

    # Header with back and edit buttons
    col1, col2, col3 = st.columns([3, 1, 1])

    with col1:
        st.title(f"{question.get_display_type()} Question")

    with col2:
        if not is_editing and st.button("✏️ Edit", width="stretch"):
            st.session_state[edit_mode_key] = True
            st.rerun()

    with col3:
        if st.button("← Back", width="stretch"):
            if is_editing:
                st.session_state[edit_mode_key] = False
            del st.session_state['view_question_id']
            st.rerun()

    # Show edit form if in edit mode
    if is_editing:
        show_question_edit_form(db, question, question_id, edit_mode_key)
        return

    st.divider()

    # Metadata row
    col1, col2, col3, col4 = st.columns(4)

    with col1:
        st.metric("Difficulty", f"{question.get_difficulty_emoji()} {question.difficulty.title()}")

    with col2:
        st.metric("Confidence", f"{question.get_confidence_emoji()} {question.confidence_level}/5")

    with col3:
        st.metric("Practice Count", question.practice_count)

    with col4:
        if question.last_practiced:
            last_practiced = question.last_practiced[:10]
            st.metric("Last Practiced", last_practiced)
        else:
            st.metric("Last Practiced", "Never")

    st.divider()

    # Question details tabs
    tab1, tab2, tab3 = st.tabs(["📝 Question & Answer", "ℹ️ Details", "⚙️ Actions"])

    with tab1:
        st.subheader("Question")
        st.markdown(f"**{question.question}**")

        st.markdown("---")

        st.subheader("Answer")
        st.markdown(question.answer_full)

        if question.answer_star:
            st.markdown("---")
            st.subheader("📋 STAR Format")

            if 'situation' in question.answer_star:
                st.markdown(f"**Situation:** {question.answer_star['situation']}")
            if 'task' in question.answer_star:
                st.markdown(f"**Task:** {question.answer_star['task']}")
            if 'action' in question.answer_star:
                st.markdown(f"**Action:** {question.answer_star['action']}")
            if 'result' in question.answer_star:
                st.markdown(f"**Result:** {question.answer_star['result']}")

        if question.notes:
            st.markdown("---")
            st.subheader("📝 Notes")
            st.markdown(question.notes)

    with tab2:
        st.subheader("Question Information")

        col1, col2 = st.columns(2)

        with col1:
            st.write("**Type:**", question.get_display_type())
            st.write("**Category:**", question.category.title())
            st.write("**Difficulty:**", f"{question.get_difficulty_emoji()} {question.difficulty.title()}")

        with col2:
            st.write("**Confidence Level:**", f"{question.get_confidence_emoji()} {question.confidence_level}/5")
            st.write("**Practice Count:**", question.practice_count)
            if question.last_practiced:
                st.write("**Last Practiced:**", question.last_practiced[:10])

        if question.tags:
            st.write("**Tags:**", ", ".join(question.tags))

        st.divider()

        st.write("**Created:**", question.created_at[:10])
        st.write("**Updated:**", question.updated_at[:10])

    with tab3:
        col1, col2 = st.columns(2)

        with col1:
            st.markdown("### Update Confidence")
            
            # Track the last saved confidence level
            confidence_track_key = f'saved_confidence_{question.id}'
            saved_confidence = st.session_state.get(confidence_track_key, question.confidence_level)
            
            new_confidence = st.slider(
                "Confidence Level",
                min_value=1,
                max_value=5,
                value=question.confidence_level,
                key=f"confidence_slider_{question.id}",
                help="How confident are you with this answer? (Auto-saves on change)"
            )
            
            # Auto-save when confidence level changes from the saved value
            if new_confidence != saved_confidence:
                question.update_confidence(new_confidence)
                db.update_question(question)
                st.session_state[confidence_track_key] = new_confidence
                st.success("✅ Confidence updated!")
                st.rerun()

        with col2:
            st.markdown("### Mark as Practiced")
            st.write(f"Current practice count: {question.practice_count}")

            if st.button("✅ Mark as Practiced", width="stretch"):
                question.mark_practiced()
                db.update_question(question)
                st.success("✅ Marked as practiced!")
                st.rerun()

        st.divider()

        # Update Answer Section
        st.markdown("### ✏️ Update Answer")

        # Control form visibility - automatically expand when section is active
        # Collapses after update
        edit_answer_key = f'edit_answer_expanded_{question.id}'
        collapse_flag_key = f'just_collapsed_answer_{question.id}'
        
        # Check if we just collapsed (from Update Answer)
        just_collapsed = st.session_state.get(collapse_flag_key, False)
        
        # If we just collapsed, keep it collapsed for this render
        # The flag will be cleared when user interacts with any widget in the section
        if just_collapsed:
            st.session_state[edit_answer_key] = False
        else:
            # Section is active and we didn't just collapse, so show the form
            st.session_state[edit_answer_key] = True
        
        # Read directly from session state (should be set above)
        edit_answer_expanded = st.session_state.get(edit_answer_key, True)
        
        # If collapsed, show a button to expand
        if not edit_answer_expanded:
            if st.button("✏️ Edit Answer", type="primary", key=f"expand_edit_answer_{question.id}"):
                st.session_state[edit_answer_key] = True
                st.session_state[collapse_flag_key] = False
                st.rerun()
        
        # Render the form
        if edit_answer_expanded:
            new_answer = st.text_area(
                "Answer",
                value=question.answer_full,
                height=200,
                help="Update your answer to this question"
            )

            # Optional: Update STAR format
            st.markdown("**STAR Format (Optional)**")

            col1, col2 = st.columns(2)

            with col1:
                star_situation = st.text_area(
                    "Situation",
                    value=question.answer_star.get('situation', '') if question.answer_star else '',
                    height=80,
                    key=f"star_situation_{question.id}"
                )

                star_task = st.text_area(
                    "Task",
                    value=question.answer_star.get('task', '') if question.answer_star else '',
                    height=80,
                    key=f"star_task_{question.id}"
                )

            with col2:
                star_action = st.text_area(
                    "Action",
                    value=question.answer_star.get('action', '') if question.answer_star else '',
                    height=80,
                    key=f"star_action_{question.id}"
                )

                star_result = st.text_area(
                    "Result",
                    value=question.answer_star.get('result', '') if question.answer_star else '',
                    height=80,
                    key=f"star_result_{question.id}"
                )

            notes_update = st.text_area(
                "Notes",
                value=question.notes,
                height=100,
                help="Additional notes about this question"
            )

            if st.button("💾 Update Answer", type="primary", width="stretch"):
                # Update answer
                question.answer_full = new_answer

                # Update STAR format if any field is filled
                if star_situation or star_task or star_action or star_result:
                    question.answer_star = {
                        'situation': star_situation,
                        'task': star_task,
                        'action': star_action,
                        'result': star_result
                    }

                # Update notes
                question.notes = notes_update

                # Update timestamp
                question.updated_at = datetime.now().isoformat()

                # Save to database
                db.update_question(question)
                
                # Collapse the form and set flag to prevent immediate re-expansion
                st.session_state[edit_answer_key] = False
                st.session_state[collapse_flag_key] = True
                
                st.success("✅ Answer updated successfully!")
                st.rerun()

        st.divider()

        # Danger zone
        with st.expander("🗑️ Danger Zone"):
            st.warning("This action cannot be undone!")

            if st.button("Delete Question", type="secondary"):
                if st.session_state.get('confirm_delete_question'):
                    db.delete_question(question_id)
                    del st.session_state['view_question_id']
                    st.success("Question deleted!")
                    st.rerun()
                else:
                    st.session_state['confirm_delete_question'] = True
                    st.warning("Click again to confirm deletion")


def show_recent_questions(db: InterviewDB, limit: int = 10):
    """Show recently added questions"""
    questions = db.list_questions()

    if not questions:
        st.info("No questions yet. Add your first one above!")
        return

    # Sort by created_at (most recent first)
    questions.sort(key=lambda x: x.created_at, reverse=True)
    questions = questions[:limit]

    st.subheader(f"📋 Recent Questions ({len(questions)})")

    for q in questions:
        with st.container():
            col1, col2, col3 = st.columns([4, 2, 1])

            with col1:
                st.markdown(f"**{q.question}**")
                # Badges
                badges = f"{q.get_display_type()} • {q.get_difficulty_emoji()} {q.difficulty.title()} • {q.category.title()}"
                st.caption(badges)

            with col2:
                st.write(f"Confidence: {q.get_confidence_emoji()} {q.confidence_level}/5")
                if q.practice_count > 0:
                    st.caption(f"Practiced {q.practice_count}x")

            with col3:
                if st.button("View", key=f"view_{q.id}"):
                    st.session_state['view_question_id'] = q.id
                    st.rerun()

            st.divider()


def login_screen():
    # Hide sidebar navigation before login and style the login button with Google blue
    st.markdown("""
        <style>
            [data-testid="stSidebar"] {
                display: none;
            }
            /* Google blue branding for login button */
            .stButton > button[kind="primary"] {
                background-color: #4285F4 !important;
                border-color: #4285F4 !important;
            }
            .stButton > button[kind="primary"]:hover {
                background-color: #357ae8 !important;
                border-color: #357ae8 !important;
            }
            .stButton > button[kind="primary"]:active {
                background-color: #2a66c9 !important;
                border-color: #2a66c9 !important;
            }
        
            /* LinkedIn blue branding for link button */
            div[data-testid="stLinkButton"] {
                width: 100% !important;
                display: block !important;
            }
            div[data-testid="stLinkButton"] > a {
                background-color: #0077B5 !important;
                color: white !important;
                border: 1px solid #0077B5 !important;
                width: 100% !important;
                text-align: center !important;
                display: block !important;
                padding: 0.5rem 0.75rem !important;
                border-radius: 0.5rem !important;
                text-decoration: none !important;
                font-weight: 400 !important;
                font-size: 1rem !important;
                line-height: 1.6 !important;
                min-height: 2.5rem !important;
                box-sizing: border-box !important;
            }
            div[data-testid="stLinkButton"] > a:hover {
                background-color: #006396 !important;
                border-color: #006396 !important;
                color: white !important;
            }
            div[data-testid="stLinkButton"] > a:active {
                background-color: #005077 !important;
                border-color: #005077 !important;
                color: white !important;
            }
        </style>
    """, unsafe_allow_html=True)

    st.header("Please log in to access Interview Prep")
    st.subheader("Please log in.")
    render_login_button(type="primary", use_container_width=True)
    
    # LinkedIn login (if configured)
    if is_linkedin_configured():
        st.markdown("<div style='margin-top: 12px;'></div>", unsafe_allow_html=True)
        render_linkedin_login_button(label="🔗 Login with LinkedIn")


def main():
    st.set_page_config(page_title="Interview Prep", page_icon="🎯", layout="wide")

    # Handle LinkedIn OAuth callback
    query_params = st.query_params

    # Check if this is a LinkedIn callback (has code and state parameters)
    if 'code' in query_params and 'state' in query_params and query_params.get('state', '').startswith('linkedin_'):
        # This is a LinkedIn OAuth callback
        code = query_params['code']
        state = query_params['state']

        # Only process if not already processed (check session state flag)
        if not st.session_state.get('linkedin_callback_processed'):
            # Show processing message
            with st.spinner("Processing LinkedIn login..."):
                # Handle the callback
                success = handle_linkedin_callback(code, state)

            if success:
                # Mark as processed to prevent re-processing on rerun
                st.session_state['linkedin_callback_processed'] = True

                # Clear query parameters
                try:
                    st.query_params.clear()
                except:
                    pass

                # Small delay to ensure session state is saved
                import time
                time.sleep(0.5)

                # Rerun to show the main app (user is now logged in)
                st.rerun()
            else:
                st.error("Failed to complete LinkedIn login. Please try again.")
                if 'linkedin_login_initiated' in st.session_state:
                    del st.session_state['linkedin_login_initiated']
                if 'linkedin_oauth_state' in st.session_state:
                    del st.session_state['linkedin_oauth_state']

                try:
                    st.query_params.clear()
                except:
                    pass

                st.rerun()
        else:
            try:
                st.query_params.clear()
            except:
                pass



    # Apply Google blue to all primary buttons
    from components.styles import apply_google_button_style
    apply_google_button_style()

    if not is_user_logged_in():
        login_screen()
        return

    st.title("🎯 Interview Preparation")
    st.markdown("Build your personal interview toolkit")

    # Initialize database
    db = InterviewDB()

    # Check if viewing a specific question
    if st.session_state.get('view_question_id'):
        show_question_detail(db, st.session_state['view_question_id'])
        return

    # Get stats
    stats = db.get_stats()

    # Key Metrics Row
    st.header("📊 Your Prep Stats")

    col1, col2, col3 = st.columns(3)

    with col1:
        st.metric(
            "Questions",
            stats['total_questions'],
            help="Total questions in your bank"
        )

    with col2:
        st.metric(
            "Practiced",
            f"{stats['practice_percentage']:.0f}%",
            help="Percentage of questions you've practiced"
        )

    with col3:
        st.metric(
            "Practice Hours",
            f"{stats['total_practice_time_hours']:.1f}",
            help="Total practice time"
        )

    st.divider()

    # Quick Actions
    st.header("⚡ Quick Actions")

    col1, col2, col3, col4 = st.columns(4)

    with col1:
        if st.button("📄 Upload Document", width="stretch"):
            st.session_state['show_upload_document'] = True
            st.session_state['show_add_question'] = False

    with col2:
        if st.button("➕ Add Question", width="stretch"):
            st.session_state['show_add_question'] = True
            st.session_state['show_upload_document'] = False

    with col3:
        if st.button("📝 View All Questions", width="stretch"):
            st.switch_page("pages/questions.py")

    with col4:
        if st.button("🎓 Practice Mode", width="stretch"):
            st.session_state['show_practice_mode'] = True
            st.session_state['show_add_question'] = False
            st.session_state['show_upload_document'] = False

    st.divider()

    # Show upload document form if requested
    if st.session_state.get('show_upload_document', False):
        show_upload_document_form(db)

        if st.button("✕ Close Upload"):
            st.session_state['show_upload_document'] = False
            st.rerun()

        st.divider()

    # Show add question form if requested
    if st.session_state.get('show_add_question', False):
        show_add_question_form(db)

        if st.button("✕ Close Form"):
            st.session_state['show_add_question'] = False
            st.rerun()

        st.divider()

    # Show parsed questions review if available
    if 'parsed_questions' in st.session_state:
        show_parsed_questions_review(db)
        st.divider()

    # Show practice mode if requested
    if st.session_state.get('show_practice_mode', False):
        show_practice_mode(db)

        if st.button("✕ Close Practice Mode"):
            st.session_state['show_practice_mode'] = False
            st.session_state['practice_session'] = None
            st.session_state['current_question_index'] = 0
            st.session_state['practice_questions'] = []
            st.session_state['start_time'] = None
            st.session_state['show_answer'] = False
            st.rerun()

        st.divider()

    # Breakdown by type
    if stats['total_questions'] > 0:
        st.header("📈 Question Breakdown")

        col1, col2 = st.columns(2)

        with col1:
            st.subheader("By Type")
            for q_type, count in stats['questions_by_type'].items():
                percentage = (count / stats['total_questions'] * 100)
                st.write(f"**{q_type.title()}:** {count} ({percentage:.0f}%)")

        with col2:
            st.subheader("By Difficulty")
            for difficulty, count in stats['questions_by_difficulty'].items():
                percentage = (count / stats['total_questions'] * 100)
                emoji = {'easy': '🟢', 'medium': '🟡', 'hard': '🔴'}.get(difficulty, '⚪')
                st.write(f"**{emoji} {difficulty.title()}:** {count} ({percentage:.0f}%)")

        st.divider()

    # Recent questions
    show_recent_questions(db)

    # Navigation buttons
    st.divider()
    st.header("🧭 Navigation")

    col1, col2, col3 = st.columns(3)

    with col1:
        if st.button("🏠 Home", width="stretch"):
            st.switch_page("app.py")

    with col2:
        if st.button("📊 Dashboard", width="stretch"):
            st.switch_page("pages/dashboard.py")

    with col3:
        if st.button("📝 Applications", width="stretch"):
            st.switch_page("pages/applications.py")
    
    # Logout button
    st.divider()
    st.button("Log out", on_click=logout)


if __name__ == "__main__":
    main()
